"""
Gerador de Documentos do Sistema Vaucher e Álvares
Migrado do main.py em 19/01/2026

Este arquivo contém a classe GeradorDocumentos e funções para geração de DOCX.
"""

import os
import uuid
import zipfile
import shutil
from datetime import datetime

from modules.config import MODELOS_DIR, GERADOS_DIR, logger


class GeradorDocumentos:
    """Classe para gerar documentos DOCX a partir de modelos."""

    def __init__(self):
        self.modelo_contrato = os.path.join(MODELOS_DIR, 'CONTRATO_Modelo.docx')
        self.modelo_procuracao = os.path.join(MODELOS_DIR, 'Procuracao_Modelo.docx')
        self.modelo_prestacao = os.path.join(MODELOS_DIR, 'Prestacao_Contas_Modelo.docx')

    def _formatar_data(self, data_str: str) -> str:
        """Formata data de YYYY-MM-DD para DD/MM/YYYY."""
        if not data_str:
            return ''
        try:
            if '-' in data_str:
                partes = data_str.split('-')
                return f"{partes[2]}/{partes[1]}/{partes[0]}"
            return data_str
        except:
            return data_str

    def _data_por_extenso(self) -> str:
        """Retorna data atual por extenso."""
        meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
                 'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
        hoje = datetime.now()
        return f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

    def _substituir_no_xml(self, xml_content: str, dados: dict) -> str:
        """Substitui placeholders no XML do documento."""
        substituicoes = {
            '{{nome}}': dados.get('nome', ''),
            '{{nacionalidade}}': dados.get('nacionalidade', ''),
            '{{estado_civil}}': dados.get('estado_civil', ''),
            '{{profissão}}': dados.get('profissao', ''),
            '{{rg}}': dados.get('rg', ''),
            '{{cpf}}': dados.get('cpf', ''),
            '{{data_nascimento}}': self._formatar_data(dados.get('data_nascimento', '')),
            '{{endereco_completo}}': dados.get('endereco_completo', ''),
            '{{email}}': dados.get('email', ''),
            '{{telefone}}': dados.get('telefone', ''),
            '{{poderes_especificos}}': dados.get('poderes_especificos', ''),
        }

        resultado = xml_content
        for placeholder, valor in substituicoes.items():
            resultado = resultado.replace(placeholder, valor)

        objeto = dados.get('objeto_contrato', '')
        if objeto:
            resultado = resultado.replace(
                'advocatícios para .',
                f'advocatícios para {objeto}.'
            )

        honorarios = dados.get('honorarios', '')
        if honorarios:
            resultado = resultado.replace(
                'fixar-se-ão em .',
                f'fixar-se-ão em {honorarios}.'
            )

        resultado = resultado.replace('sample text question answer', self._data_por_extenso())
        resultado = resultado.replace(
            'Cuiabá, ____ de ____________de________.',
            f'Cuiabá, {self._data_por_extenso()}.'
        )

        return resultado

    def _gerar_documento(self, modelo_path: str, dados: dict, nome_saida: str, cadastro_id: str) -> str:
        """Gera documento a partir de modelo DOCX."""
        if not os.path.exists(modelo_path):
            raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

        cliente_dir = os.path.join(GERADOS_DIR, cadastro_id)
        os.makedirs(cliente_dir, exist_ok=True)

        temp_dir = os.path.join(cliente_dir, f'temp_{uuid.uuid4().hex[:8]}')
        os.makedirs(temp_dir, exist_ok=True)

        try:
            with zipfile.ZipFile(modelo_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                content = f.read()

            content = self._substituir_no_xml(content, dados)

            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(content)

            saida_path = os.path.join(cliente_dir, nome_saida)

            with zipfile.ZipFile(saida_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, temp_dir)
                        zipf.write(file_path, arcname)

            return saida_path
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)

    def gerar_contrato(self, dados: dict, cadastro_id: str) -> str:
        """Gera contrato de honorários."""
        nome = dados.get('nome', 'Cliente').replace(' ', '_')
        nome_arquivo = f"Contrato_Honorarios_{nome}.docx"
        return self._gerar_documento(self.modelo_contrato, dados, nome_arquivo, cadastro_id)

    def gerar_procuracao(self, dados: dict, cadastro_id: str) -> str:
        """Gera procuração."""
        nome = dados.get('nome', 'Cliente').replace(' ', '_')
        nome_arquivo = f"Procuracao_{nome}.docx"
        return self._gerar_documento(self.modelo_procuracao, dados, nome_arquivo, cadastro_id)

    def _format_money(self, value: float) -> str:
        """Formata valor para moeda brasileira."""
        return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def _format_data(self, data_str: str) -> str:
        """Formata data de YYYY-MM-DD para DD/MM/YYYY."""
        if not data_str:
            return ""
        try:
            partes = data_str.split("-")
            return f"{partes[2]}/{partes[1]}/{partes[0]}"
        except:
            return data_str

    def gerar_prestacao_contas(self, dados_cliente: dict, financeiro: dict, cadastro_id: str) -> str:
        """Gera documento de prestação de contas completo usando python-docx."""
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        depositos = financeiro.get("depositos", [])
        sucumbencias = financeiro.get("sucumbencias", [])
        retencoes = financeiro.get("retencoes", [])
        percentual = float(financeiro.get("percentual_honorarios", 20))
        valor_credito = float(financeiro.get("valor_credito_cliente", 0))

        total_depositos = sum(float(d.get("valor", 0)) for d in depositos)
        total_sucumbencias = sum(float(s.get("valor", 0)) for s in sucumbencias)
        total_retencoes = sum(float(r.get("valor", 0)) for r in retencoes)
        honorarios_contratuais = valor_credito * (percentual / 100)
        valor_liquido = valor_credito - honorarios_contratuais - total_retencoes

        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        def set_cell_shading(cell, color):
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            cell._tc.get_or_add_tcPr().append(shading)

        # TÍTULO
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("PRESTAÇÃO DE CONTAS")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Arial"

        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = subtitulo.add_run("VAUCHER E ÁLVARES SOCIEDADE DE ADVOGADOS ")
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = subtitulo.add_run("→ ")
        run2.font.size = Pt(11)
        run3 = subtitulo.add_run("Cliente")
        run3.bold = True
        run3.font.size = Pt(11)
        run3.font.color.rgb = RGBColor(238, 0, 0)

        # 1. IDENTIFICAÇÃO
        doc.add_paragraph()
        h1 = doc.add_paragraph()
        run = h1.add_run("1. Identificação das Partes")
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        p.add_run("Cliente: ").bold = True
        p.add_run(dados_cliente.get('nome', '').upper())

        p = doc.add_paragraph()
        p.add_run("Escritório de Advocacia: ").bold = True
        p.add_run("VAUCHER E ÁLVARES SOCIEDADE DE ADVOGADOS").bold = True
        p.add_run(", devidamente registrada na Ordem dos Advogados do Brasil Seccional de Mato Grosso sob o nº 669, inscrita no CNPJ sob o nº 21.336.697/0001-46, com sede na Rua Lima, n. 106, Bairro Jardim das Américas, em Cuiabá-MT.")

        p = doc.add_paragraph()
        p.add_run("Processo(s): ").bold = True
        run = p.add_run(f"{financeiro.get('numero_processo', '')} / {financeiro.get('vara_tribunal', '')}")
        run.font.color.rgb = RGBColor(238, 0, 0)

        # 2. OBJETO
        doc.add_paragraph()
        h2 = doc.add_paragraph()
        run = h2.add_run("2. Objeto da Prestação de Contas")
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.add_run("A presente prestação de contas tem por finalidade demonstrar, de forma ")
        p.add_run("transparente, discriminada e fundamentada").bold = True
        p.add_run(", os valores ")
        p.add_run("totais recebidos").bold = True
        p.add_run(" no âmbito do(s) processo(s) acima identificado(s), indicando:")

        items = [
            ("valores pertencentes ao ", "cliente", ";"),
            ("valores correspondentes aos ", "honorários advocatícios contratuais", ";"),
            ("valores referentes aos ", "honorários advocatícios sucumbenciais", ", de titularidade do advogado;"),
            ("valores retidos a título de ", "tributos/contribuição previdenciária (PSS)", "."),
        ]
        for prefix, bold_text, suffix in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(prefix)
            p.add_run(bold_text).bold = True
            p.add_run(suffix)

        # 3. VALORES TOTAIS RECEBIDOS
        doc.add_paragraph()
        h3 = doc.add_paragraph()
        run = h3.add_run("3. Valores Totais Recebidos")
        run.bold = True
        run.font.size = Pt(11)

        num_depositos = len(depositos) if depositos else 1
        table = doc.add_table(rows=num_depositos + 2, cols=3)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        hdr[0].text = "Data do Recebimento"
        hdr[1].text = "Origem do Valor"
        hdr[2].text = "Valor Bruto (R$)"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, "D9D9D9")

        if depositos:
            for i, dep in enumerate(depositos):
                row = table.rows[i + 1].cells
                row[0].text = self._format_data(dep.get("data", ""))
                row[1].text = dep.get("origem", "")
                row[2].text = self._format_money(float(dep.get("valor", 0)))
        else:
            row = table.rows[1].cells
            row[0].text = "-"
            row[1].text = "-"
            row[2].text = self._format_money(0)

        total_row = table.rows[-1].cells
        total_row[0].text = ""
        total_row[1].text = "TOTAL"
        total_row[1].paragraphs[0].runs[0].bold = True
        total_row[2].text = self._format_money(total_depositos)
        total_row[2].paragraphs[0].runs[0].bold = True

        # 4. DISCRIMINAÇÃO DOS VALORES
        doc.add_paragraph()
        h4 = doc.add_paragraph()
        run = h4.add_run("4. Discriminação dos Valores")
        run.bold = True
        run.font.size = Pt(11)

        # 4.1 Receita do Cliente
        h41 = doc.add_paragraph()
        run = h41.add_run("4.1. Receita Pertencente ao Cliente")
        run.bold = True

        p = doc.add_paragraph()
        p.add_run("Corresponde à parcela do valor recebido que integra o patrimônio do cliente, após a dedução dos honorários advocatícios devidos e das retenções legais.")

        table41 = doc.add_table(rows=5 + len(retencoes), cols=2)
        table41.style = 'Table Grid'

        rows_data = [
            ("Valor bruto total recebido (Principal + Sucumbência)", self._format_money(total_depositos), False),
            (f"(-) Honorários contratuais ({percentual}% sobre crédito do cliente)", self._format_money(honorarios_contratuais), False),
            (f"(-) Honorários sucumbenciais", self._format_money(total_sucumbencias), False),
        ]

        for ret in retencoes:
            rows_data.append((f"(-) {ret.get('descricao', 'Retenção')}", self._format_money(float(ret.get('valor', 0))), False))

        if not retencoes:
            rows_data.append(("(-) Retenções Legais (PSS/IRRF)", self._format_money(0), False))

        rows_data.append(("Valor líquido devido ao cliente", self._format_money(valor_liquido), True))

        for i, (desc, val, is_total) in enumerate(rows_data):
            row = table41.rows[i].cells
            row[0].text = desc
            row[1].text = val
            if is_total:
                row[0].paragraphs[0].runs[0].bold = True
                row[1].paragraphs[0].runs[0].bold = True
                set_cell_shading(row[0], "E2EFDA")
                set_cell_shading(row[1], "E2EFDA")

        # 4.2 Honorários Contratuais
        doc.add_paragraph()
        h42 = doc.add_paragraph()
        run = h42.add_run("4.2. Honorários Advocatícios Contratuais")
        run.bold = True

        p = doc.add_paragraph()
        p.add_run("Nos termos do art. 22 da Lei nº 8.906/1994, os honorários advocatícios ajustados em contrato constituem direito do advogado, possuindo natureza remuneratória pelos serviços prestados.")

        p = doc.add_paragraph()
        p.add_run("Percentual contratado: ").bold = True
        p.add_run(f"{percentual}%")

        p = doc.add_paragraph()
        p.add_run("Base de cálculo: ").bold = True
        p.add_run(f"Valor do crédito do cliente ({self._format_money(valor_credito)})")

        table42 = doc.add_table(rows=1, cols=2)
        table42.style = 'Table Grid'
        row = table42.rows[0].cells
        row[0].text = f"Percentual contratual ({percentual}%) sobre {self._format_money(valor_credito)}"
        row[1].text = self._format_money(honorarios_contratuais)
        row[1].paragraphs[0].runs[0].bold = True

        # 4.3 Honorários Sucumbenciais
        doc.add_paragraph()
        h43 = doc.add_paragraph()
        run = h43.add_run("4.3. Honorários Advocatícios Sucumbenciais")
        run.bold = True

        p = doc.add_paragraph()
        p.add_run("Os honorários sucumbenciais são fixados judicialmente e pertencem exclusivamente ao advogado, conforme dispõe expressamente o art. 85, §14, do CPC.")

        num_sucumb = len(sucumbencias) if sucumbencias else 1
        table43 = doc.add_table(rows=num_sucumb + 1, cols=2)
        table43.style = 'Table Grid'

        if sucumbencias:
            for i, suc in enumerate(sucumbencias):
                row = table43.rows[i].cells
                row[0].text = suc.get("descricao", "")
                row[1].text = self._format_money(float(suc.get("valor", 0)))
        else:
            row = table43.rows[0].cells
            row[0].text = "Honorários sucumbenciais"
            row[1].text = self._format_money(0)

        total_suc_row = table43.rows[-1].cells
        total_suc_row[0].text = "Total Honorários Sucumbenciais"
        total_suc_row[0].paragraphs[0].runs[0].bold = True
        total_suc_row[1].text = self._format_money(total_sucumbencias)
        total_suc_row[1].paragraphs[0].runs[0].bold = True

        p = doc.add_paragraph()
        run = p.add_run("Obs.: Os honorários sucumbenciais não se confundem com o crédito do cliente, não integram sua base patrimonial e não substituem os honorários contratuais.")
        run.italic = True
        run.font.size = Pt(9)

        # 5. RESUMO GERAL
        doc.add_paragraph()
        h5 = doc.add_paragraph()
        run = h5.add_run("5. Resumo Geral da Prestação de Contas")
        run.bold = True
        run.font.size = Pt(11)

        table5 = doc.add_table(rows=5, cols=3)
        table5.style = 'Table Grid'

        hdr5 = table5.rows[0].cells
        hdr5[0].text = "Natureza do Valor"
        hdr5[1].text = "Valor (R$)"
        hdr5[2].text = "Titularidade"
        for cell in hdr5:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, "D9D9D9")

        resumo_data = [
            ("Receita líquida do cliente", self._format_money(valor_liquido), "Cliente"),
            ("Honorários contratuais", self._format_money(honorarios_contratuais), "Escritório"),
            ("Honorários sucumbenciais", self._format_money(total_sucumbencias), "Escritório"),
            ("TOTAL GERAL", self._format_money(total_depositos), ""),
        ]

        for i, (nat, val, tit) in enumerate(resumo_data):
            row = table5.rows[i + 1].cells
            row[0].text = nat
            row[1].text = val
            row[2].text = tit
            if i == 3:
                row[0].paragraphs[0].runs[0].bold = True
                row[1].paragraphs[0].runs[0].bold = True
                set_cell_shading(row[0], "F2F2F2")
                set_cell_shading(row[1], "F2F2F2")
                set_cell_shading(row[2], "F2F2F2")

        # 6. CONCLUSÃO
        doc.add_paragraph()
        h6 = doc.add_paragraph()
        run = h6.add_run("6. Conclusão")
        run.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph("O escritório declara que:")

        conclusoes = [
            "os valores foram corretamente recebidos e contabilizados;",
            "a retenção dos honorários observa expressa previsão legal e contratual;",
            "o valor líquido indicado encontra-se à disposição do cliente, após a assinatura da presente prestação de contas que também reconhece a quitação geral e irrestrita quanto as obrigações do escritório na demanda em referência."
        ]
        for c in conclusoes:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(c)

        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph(f"Cuiabá-MT, {self._data_por_extenso()}.")

        doc.add_paragraph()
        doc.add_paragraph()

        ass1 = doc.add_paragraph("VAUCHER E ÁLVARES SOCIEDADE DE ADVOGADOS")
        ass1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ass1.runs[0].bold = True

        cnpj = doc.add_paragraph("CNPJ 21.336.697/0001-46")
        cnpj.alignment = WD_ALIGN_PARAGRAPH.CENTER

        oab = doc.add_paragraph("OAB/MT 669")
        oab.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        ass2 = doc.add_paragraph(dados_cliente.get('nome', '').upper())
        ass2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ass2.runs[0].bold = True

        cliente_dir = os.path.join(GERADOS_DIR, cadastro_id)
        os.makedirs(cliente_dir, exist_ok=True)

        nome = dados_cliente.get('nome', 'Cliente').replace(' ', '_')
        nome_arquivo = f"Prestacao_Contas_{nome}.docx"
        caminho_arquivo = os.path.join(cliente_dir, nome_arquivo)

        doc.save(caminho_arquivo)
        return caminho_arquivo

    def gerar_todos(self, dados: dict, cadastro_id: str) -> dict:
        """Gera todos os documentos padrão (contrato e procuração)."""
        return {
            'contrato': self.gerar_contrato(dados, cadastro_id),
            'procuracao': self.gerar_procuracao(dados, cadastro_id)
        }


# Instância global do gerador
gerador = GeradorDocumentos()


# ============================================
# FUNÇÃO PARA GERAR PETIÇÃO - AUXÍLIO MORADIA
# ============================================

def gerar_peticao_auxilio_moradia(dados_cliente: dict, dados_residencia: dict, cadastro_id: str) -> str:
    """
    Gera a petição inicial de auxílio moradia para residência médica.
    Usa o modelo peticao_auxilio_moradia_modelo.docx com substituição de placeholders.
    """

    # Caminho do modelo
    modelo_path = os.path.join(MODELOS_DIR, 'peticao_auxilio_moradia_modelo.docx')

    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

    # Funções auxiliares
    def formatar_data(data_str):
        """Formata data de YYYY-MM-DD para DD/MM/YYYY."""
        if not data_str:
            return ''
        try:
            partes = data_str.split('-')
            return f"{partes[2]}/{partes[1]}/{partes[0]}"
        except:
            return data_str

    def formatar_moeda(valor):
        """Formata valor para moeda brasileira."""
        try:
            valor_float = float(valor)
            return f"R$ {valor_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            return "R$ 0,00"

    def valor_por_extenso(valor):
        """Converte valor para extenso (simplificado)."""
        try:
            valor_float = float(valor)
            inteiro = int(valor_float)
            centavos = int((valor_float - inteiro) * 100)
            if centavos > 0:
                return f"{inteiro} reais e {centavos} centavos"
            return f"{inteiro} reais"
        except:
            return "zero reais"

    # Extrair dados do cliente
    nome = dados_cliente.get('nome', '').upper()
    cpf = dados_cliente.get('cpf', '')
    endereco = dados_cliente.get('endereco_completo', '')
    estado_civil = dados_cliente.get('estado_civil', '')

    # Extrair dados da residência médica
    hospital_nome = dados_residencia.get('unidade_hospitalar', '')
    hospital_nome_completo = dados_residencia.get('unidade_hospitalar', '')
    cnpj_hospital = dados_residencia.get('cnpj_hospital', '')
    endereco_hospital = dados_residencia.get('endereco_hospital', '')

    universidade_nome = dados_residencia.get('instituicao_ensino', '')
    cnpj_universidade = dados_residencia.get('cnpj_universidade', '')
    endereco_universidade = dados_residencia.get('endereco_universidade', '')

    especialidade = dados_residencia.get('especialidade_medica', '')
    data_inicio = dados_residencia.get('data_inicio_residencia', '')
    data_termino = dados_residencia.get('data_termino_residencia', '')
    valor_bolsa = dados_residencia.get('valor_bolsa_mensal', 0)

    # Dados do processo anterior (se houver)
    numero_processo_anterior_1 = dados_residencia.get('numero_processo_anterior', '')
    numero_processo_anterior_2 = dados_residencia.get('numero_processo_anterior_2', '')
    vara_anterior_1 = dados_residencia.get('vara_juizado_anterior', '')
    vara_anterior_2 = dados_residencia.get('vara_juizado_anterior_2', '')
    data_ajuizamento_anterior = dados_residencia.get('data_protocolo_anterior', '')

    # Calcular valores
    try:
        valor_bolsa_float = float(valor_bolsa) if valor_bolsa else 0
        valor_auxilio_moradia = valor_bolsa_float * 0.30

        # Calcular período e total
        if data_inicio and data_termino:
            from datetime import datetime as dt
            dt_inicio = dt.strptime(data_inicio, '%Y-%m-%d')
            dt_termino = dt.strptime(data_termino, '%Y-%m-%d')
            meses = (dt_termino.year - dt_inicio.year) * 12 + (dt_termino.month - dt_inicio.month)
            if meses < 0:
                meses = 0
        else:
            meses = 0

        valor_total_bolsas = valor_auxilio_moradia * meses
        periodo_auxilio = f"{meses} meses"
    except:
        valor_auxilio_moradia = 0
        valor_total_bolsas = 0
        periodo_auxilio = ""

    # Determinar gênero (baseado no estado civil)
    genero_feminino = estado_civil.lower().endswith('a') if estado_civil else False
    a_o = "a" if genero_feminino else "o"
    a_vazio = "a" if genero_feminino else ""

    # Mapeamento de placeholders para valores
    substituicoes = {
        '{{NOME_AUTOR}}': nome,
        '{{ESTADO_CIVIL}}': estado_civil,
        '{{CPF}}': cpf,
        '{{ENDERECO_COMPLETO}}': endereco,
        '{{A_O}}': a_o,
        '{{A_VAZIO}}': a_vazio,
        '{{HOSPITAL_NOME_COMPLETO}}': hospital_nome_completo,
        '{{HOSPITAL_NOME}}': hospital_nome,
        '{{CNPJ_HOSPITAL}}': cnpj_hospital,
        '{{ENDERECO_HOSPITAL}}': endereco_hospital,
        '{{UNIVERSIDADE_NOME}}': universidade_nome,
        '{{CNPJ_UNIVERSIDADE}}': cnpj_universidade,
        '{{ENDERECO_UNIVERSIDADE}}': endereco_universidade,
        '{{ESPECIALIDADE}}': especialidade,
        '{{DATA_INICIO_RESIDENCIA}}': formatar_data(data_inicio),
        '{{DATA_FIM_RESIDENCIA}}': formatar_data(data_termino),
        '{{PERIODO_AUXILIO}}': periodo_auxilio,
        '{{VALOR_BOLSA_MENSAL}}': formatar_moeda(valor_bolsa),
        '{{VALOR_AUXILIO_MORADIA}}': formatar_moeda(valor_auxilio_moradia),
        '{{VALOR_AUXILIO_MORADIA_EXTENSO}}': valor_por_extenso(valor_auxilio_moradia),
        '{{VALOR_TOTAL_BOLSAS}}': formatar_moeda(valor_total_bolsas),
        '{{NUMERO_PROCESSO_ANTERIOR_1}}': numero_processo_anterior_1,
        '{{NUMERO_PROCESSO_ANTERIOR_2}}': numero_processo_anterior_2,
        '{{VARA_ANTERIOR_1}}': vara_anterior_1,
        '{{VARA_ANTERIOR_2}}': vara_anterior_2,
        '{{DATA_AJUIZAMENTO_ANTERIOR}}': formatar_data(data_ajuizamento_anterior),
    }

    # Criar diretório de saída
    cliente_dir = os.path.join(GERADOS_DIR, cadastro_id)
    os.makedirs(cliente_dir, exist_ok=True)

    # Criar diretório temporário para extração
    temp_dir = os.path.join(cliente_dir, f'temp_{uuid.uuid4().hex[:8]}')
    os.makedirs(temp_dir, exist_ok=True)

    try:
        # Extrair o modelo (DOCX é um ZIP)
        with zipfile.ZipFile(modelo_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # Ler o document.xml
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Substituir todos os placeholders
        for placeholder, valor in substituicoes.items():
            content = content.replace(placeholder, str(valor) if valor else '')

        # Salvar o XML modificado
        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(content)

        # Criar o documento final
        nome_arquivo = f"Peticao_Auxilio_Moradia_{nome.replace(' ', '_')}.docx"
        caminho_arquivo = os.path.join(cliente_dir, nome_arquivo)

        # Recriar o DOCX
        with zipfile.ZipFile(caminho_arquivo, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)

        return caminho_arquivo

    finally:
        # Limpar diretório temporário
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
