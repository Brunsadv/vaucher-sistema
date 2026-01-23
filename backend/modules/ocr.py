"""
Módulo de OCR usando Google Cloud Vision API.
Processa documentos (imagens, PDFs) e extrai texto/tabelas.
Exporta para CSV (tabelas) ou DOCX (texto) e gera resumos.
"""

import os
import io
import csv
import json
import re
import logging
from typing import Optional, Tuple, List, Dict, Any
from datetime import datetime
from pathlib import Path

from google.cloud import vision
from google.oauth2 import service_account
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

logger = logging.getLogger(__name__)

# Diretórios
BASE_DIR = Path(__file__).parent.parent
UPLOADS_DIR = BASE_DIR / "uploads"
OCR_OUTPUT_DIR = UPLOADS_DIR / "ocr_output"
OCR_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


class GoogleVisionOCR:
    """Classe para processamento de OCR usando Google Cloud Vision."""

    def __init__(self):
        """Inicializa o cliente do Google Vision."""
        self.client = None
        self._init_client()

    def _init_client(self):
        """Inicializa o cliente com credenciais."""
        try:
            # Tenta usar credenciais do arquivo JSON
            credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
            if credentials_path and os.path.exists(credentials_path):
                credentials = service_account.Credentials.from_service_account_file(
                    credentials_path,
                    scopes=["https://www.googleapis.com/auth/cloud-vision"]
                )
                self.client = vision.ImageAnnotatorClient(credentials=credentials)
            else:
                # Tenta usar credenciais padrão do ambiente
                self.client = vision.ImageAnnotatorClient()
            logger.info("Google Vision client inicializado com sucesso")
        except Exception as e:
            logger.error(f"Erro ao inicializar Google Vision client: {e}")
            self.client = None

    def process_image(self, image_path: str) -> Dict[str, Any]:
        """
        Processa uma imagem e extrai texto usando OCR.

        Args:
            image_path: Caminho para o arquivo de imagem

        Returns:
            Dict com texto extraído, blocos de texto e informações de layout
        """
        if not self.client:
            raise Exception("Google Vision client não inicializado. Verifique as credenciais.")

        with io.open(image_path, 'rb') as image_file:
            content = image_file.read()

        image = vision.Image(content=content)

        # Detecta texto com layout completo
        response = self.client.document_text_detection(image=image)

        if response.error.message:
            raise Exception(f"Erro do Google Vision: {response.error.message}")

        result = {
            "texto_completo": "",
            "blocos": [],
            "paragrafos": [],
            "tabelas": [],
            "tipo_detectado": "texto",  # texto, tabela, misto
            "confianca_media": 0.0
        }

        if response.full_text_annotation:
            result["texto_completo"] = response.full_text_annotation.text

            # Processa blocos e parágrafos
            confiancas = []
            for page in response.full_text_annotation.pages:
                for block in page.blocks:
                    block_text = ""
                    block_paragraphs = []

                    for paragraph in block.paragraphs:
                        para_text = ""
                        for word in paragraph.words:
                            word_text = "".join([symbol.text for symbol in word.symbols])
                            para_text += word_text + " "
                            if word.confidence:
                                confiancas.append(word.confidence)

                        block_text += para_text.strip() + "\n"
                        block_paragraphs.append(para_text.strip())

                    result["blocos"].append({
                        "texto": block_text.strip(),
                        "tipo": self._get_block_type(block),
                        "vertices": self._get_vertices(block.bounding_box)
                    })
                    result["paragrafos"].extend(block_paragraphs)

            if confiancas:
                result["confianca_media"] = sum(confiancas) / len(confiancas)

        # Detecta se há tabelas
        result["tabelas"] = self._detect_tables(result["texto_completo"], result["blocos"])
        if result["tabelas"]:
            result["tipo_detectado"] = "tabela" if len(result["tabelas"]) > 0 else "misto"

        return result

    def _get_block_type(self, block) -> str:
        """Retorna o tipo do bloco."""
        block_type_map = {
            1: "texto",
            2: "tabela",
            3: "imagem",
            4: "regua",
            5: "codigo_barras"
        }
        return block_type_map.get(block.block_type, "desconhecido")

    def _get_vertices(self, bounding_box) -> List[Dict[str, int]]:
        """Extrai vértices do bounding box."""
        return [{"x": v.x, "y": v.y} for v in bounding_box.vertices]

    def _detect_tables(self, texto: str, blocos: List[Dict]) -> List[Dict]:
        """
        Detecta e extrai tabelas do texto.
        Usa heurísticas para identificar estruturas tabulares.
        """
        tabelas = []
        linhas = texto.split("\n")

        # Heurística 1: Linhas com múltiplos separadores (|, tabs, múltiplos espaços)
        tabela_atual = []
        for linha in linhas:
            # Detecta se a linha parece ser parte de uma tabela
            if self._is_table_row(linha):
                cells = self._split_table_row(linha)
                if cells:
                    tabela_atual.append(cells)
            else:
                if len(tabela_atual) >= 2:  # Mínimo 2 linhas para ser tabela
                    tabelas.append({
                        "linhas": tabela_atual,
                        "num_colunas": max(len(row) for row in tabela_atual),
                        "num_linhas": len(tabela_atual)
                    })
                tabela_atual = []

        # Última tabela
        if len(tabela_atual) >= 2:
            tabelas.append({
                "linhas": tabela_atual,
                "num_colunas": max(len(row) for row in tabela_atual),
                "num_linhas": len(tabela_atual)
            })

        return tabelas

    def _is_table_row(self, linha: str) -> bool:
        """Verifica se uma linha parece ser parte de uma tabela."""
        linha = linha.strip()
        if not linha:
            return False

        # Detecta separadores comuns
        if "|" in linha:
            return True
        if "\t" in linha:
            return True
        # Múltiplos espaços consecutivos (3+) indicam colunas
        if re.search(r'\s{3,}', linha):
            return True
        # Padrão de valores monetários ou números separados
        if re.search(r'R\$\s*[\d.,]+.*R\$\s*[\d.,]+', linha):
            return True
        if re.search(r'\d+[.,]\d+\s+\d+[.,]\d+', linha):
            return True

        return False

    def _split_table_row(self, linha: str) -> List[str]:
        """Divide uma linha de tabela em células."""
        linha = linha.strip()

        # Tenta diferentes separadores
        if "|" in linha:
            cells = [c.strip() for c in linha.split("|") if c.strip()]
        elif "\t" in linha:
            cells = [c.strip() for c in linha.split("\t") if c.strip()]
        else:
            # Múltiplos espaços
            cells = [c.strip() for c in re.split(r'\s{3,}', linha) if c.strip()]

        return cells if len(cells) > 1 else []

    def generate_summary(self, texto: str, max_length: int = 500) -> str:
        """
        Gera um resumo/sumário do conteúdo extraído.

        Args:
            texto: Texto completo extraído
            max_length: Tamanho máximo do resumo

        Returns:
            Resumo do documento
        """
        if not texto:
            return "Documento vazio ou sem texto detectável."

        linhas = texto.strip().split("\n")
        linhas = [l.strip() for l in linhas if l.strip()]

        if not linhas:
            return "Documento vazio ou sem texto detectável."

        resumo_partes = []

        # Identifica tipo de documento
        tipo_doc = self._identify_document_type(texto)
        resumo_partes.append(f"Tipo de documento: {tipo_doc}")

        # Extrai informações-chave
        info_chave = self._extract_key_info(texto)
        if info_chave:
            resumo_partes.append("\nInformações identificadas:")
            for chave, valor in info_chave.items():
                resumo_partes.append(f"  - {chave}: {valor}")

        # Primeiras linhas significativas
        linhas_significativas = [l for l in linhas[:10] if len(l) > 10]
        if linhas_significativas:
            resumo_partes.append("\nConteúdo principal:")
            for linha in linhas_significativas[:5]:
                if len(linha) > 100:
                    linha = linha[:100] + "..."
                resumo_partes.append(f"  {linha}")

        # Estatísticas
        num_palavras = len(texto.split())
        num_linhas = len(linhas)
        resumo_partes.append(f"\nEstatísticas: {num_palavras} palavras, {num_linhas} linhas")

        resumo = "\n".join(resumo_partes)

        if len(resumo) > max_length:
            resumo = resumo[:max_length-3] + "..."

        return resumo

    def _identify_document_type(self, texto: str) -> str:
        """Identifica o tipo de documento baseado no conteúdo."""
        texto_lower = texto.lower()

        # Comprovantes de pagamento
        if any(p in texto_lower for p in ["comprovante", "transferência", "pix", "ted", "doc", "pagamento"]):
            if "pix" in texto_lower:
                return "Comprovante de PIX"
            if any(p in texto_lower for p in ["ted", "doc", "transferência"]):
                return "Comprovante de Transferência Bancária"
            return "Comprovante de Pagamento"

        # Documentos de identidade
        if any(p in texto_lower for p in ["registro geral", "identidade", "rg"]):
            return "Documento de Identidade (RG)"
        if any(p in texto_lower for p in ["cpf", "cadastro de pessoa"]):
            return "CPF"
        if any(p in texto_lower for p in ["cnh", "carteira nacional", "habilitação"]):
            return "CNH"

        # Documentos jurídicos
        if any(p in texto_lower for p in ["procuração", "outorgante", "outorgado"]):
            return "Procuração"
        if any(p in texto_lower for p in ["contrato", "cláusula", "contratante", "contratado"]):
            return "Contrato"
        if any(p in texto_lower for p in ["petição", "excelentíssimo", "meritíssimo"]):
            return "Petição"

        # Comprovantes de residência
        if any(p in texto_lower for p in ["conta de luz", "energia elétrica", "cemig", "cpfl"]):
            return "Conta de Energia"
        if any(p in texto_lower for p in ["conta de água", "saneamento", "sabesp", "copasa"]):
            return "Conta de Água"
        if any(p in texto_lower for p in ["fatura", "telefone", "internet"]):
            return "Fatura de Serviços"

        # Outros
        if any(p in texto_lower for p in ["nota fiscal", "nf-e", "danfe"]):
            return "Nota Fiscal"
        if any(p in texto_lower for p in ["recibo", "recebi de"]):
            return "Recibo"
        if any(p in texto_lower for p in ["extrato", "saldo"]):
            return "Extrato Bancário"

        return "Documento Geral"

    def _extract_key_info(self, texto: str) -> Dict[str, str]:
        """Extrai informações-chave do documento."""
        info = {}

        # CPF
        cpf_match = re.search(r'\d{3}[.\s]?\d{3}[.\s]?\d{3}[-.\s]?\d{2}', texto)
        if cpf_match:
            info["CPF"] = cpf_match.group()

        # CNPJ
        cnpj_match = re.search(r'\d{2}[.\s]?\d{3}[.\s]?\d{3}[/.\s]?\d{4}[-.\s]?\d{2}', texto)
        if cnpj_match:
            info["CNPJ"] = cnpj_match.group()

        # Valores monetários
        valores = re.findall(r'R\$\s*[\d.,]+', texto)
        if valores:
            # Pega o maior valor como principal
            valores_num = []
            for v in valores:
                try:
                    num = float(v.replace("R$", "").replace(".", "").replace(",", ".").strip())
                    valores_num.append((v, num))
                except:
                    pass
            if valores_num:
                maior = max(valores_num, key=lambda x: x[1])
                info["Valor Principal"] = maior[0]

        # Datas
        datas = re.findall(r'\d{2}[/.-]\d{2}[/.-]\d{2,4}', texto)
        if datas:
            info["Data"] = datas[0]

        # Nomes (linhas que parecem nomes próprios)
        linhas = texto.split("\n")
        for linha in linhas[:20]:
            linha = linha.strip()
            # Nome próprio: 2-5 palavras começando com maiúscula
            if re.match(r'^[A-Z][a-záéíóúãõâêîôûç]+(\s+[A-Z][a-záéíóúãõâêîôûç]+){1,4}$', linha):
                if "Nome" not in info:
                    info["Nome"] = linha
                    break

        return info


class DocumentExporter:
    """Classe para exportar resultados de OCR para diferentes formatos."""

    def __init__(self, output_dir: Path = OCR_OUTPUT_DIR):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_to_csv(self, tabelas: List[Dict], filename: str) -> str:
        """
        Exporta tabelas detectadas para arquivo CSV.

        Args:
            tabelas: Lista de tabelas extraídas
            filename: Nome base do arquivo

        Returns:
            Caminho do arquivo gerado
        """
        if not tabelas:
            raise ValueError("Nenhuma tabela para exportar")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = self.output_dir / f"{filename}_{timestamp}.csv"

        with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f, delimiter=';')

            for i, tabela in enumerate(tabelas):
                if i > 0:
                    writer.writerow([])  # Linha em branco entre tabelas
                    writer.writerow([f"--- Tabela {i+1} ---"])

                for linha in tabela["linhas"]:
                    writer.writerow(linha)

        logger.info(f"CSV exportado: {filepath}")
        return str(filepath)

    def export_to_xlsx(self, tabelas: List[Dict], filename: str) -> str:
        """
        Exporta tabelas detectadas para arquivo Excel.

        Args:
            tabelas: Lista de tabelas extraídas
            filename: Nome base do arquivo

        Returns:
            Caminho do arquivo gerado
        """
        if not tabelas:
            raise ValueError("Nenhuma tabela para exportar")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = self.output_dir / f"{filename}_{timestamp}.xlsx"

        wb = openpyxl.Workbook()

        for i, tabela in enumerate(tabelas):
            if i == 0:
                ws = wb.active
                ws.title = f"Tabela {i+1}"
            else:
                ws = wb.create_sheet(title=f"Tabela {i+1}")

            # Estilos
            header_font = Font(bold=True)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            for row_idx, linha in enumerate(tabela["linhas"], 1):
                for col_idx, celula in enumerate(linha, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=celula)
                    cell.border = border
                    if row_idx == 1:
                        cell.font = header_font
                        cell.alignment = Alignment(horizontal='center')

            # Ajusta largura das colunas
            for col in ws.columns:
                max_length = 0
                for cell in col:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

        wb.save(filepath)
        logger.info(f"Excel exportado: {filepath}")
        return str(filepath)

    def export_to_docx(self, resultado_ocr: Dict, filename: str, incluir_resumo: bool = True) -> str:
        """
        Exporta texto extraído para arquivo DOCX.

        Args:
            resultado_ocr: Resultado do processamento OCR
            filename: Nome base do arquivo
            incluir_resumo: Se deve incluir o resumo no início

        Returns:
            Caminho do arquivo gerado
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = self.output_dir / f"{filename}_{timestamp}.docx"

        doc = Document()

        # Título
        titulo = doc.add_heading("Documento Processado por OCR", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadados
        doc.add_paragraph(f"Data de processamento: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Tipo detectado: {resultado_ocr.get('tipo_detectado', 'N/A')}")
        doc.add_paragraph(f"Confiança média: {resultado_ocr.get('confianca_media', 0):.1%}")

        # Resumo
        if incluir_resumo and resultado_ocr.get("resumo"):
            doc.add_heading("Resumo", level=1)
            doc.add_paragraph(resultado_ocr["resumo"])

        # Texto completo
        doc.add_heading("Texto Extraído", level=1)
        texto = resultado_ocr.get("texto_completo", "")

        # Divide em parágrafos
        paragrafos = texto.split("\n\n")
        for paragrafo in paragrafos:
            if paragrafo.strip():
                p = doc.add_paragraph(paragrafo.strip())
                p.paragraph_format.space_after = Pt(12)

        # Tabelas detectadas
        if resultado_ocr.get("tabelas"):
            doc.add_heading("Tabelas Detectadas", level=1)
            for i, tabela in enumerate(resultado_ocr["tabelas"], 1):
                doc.add_heading(f"Tabela {i}", level=2)

                # Cria tabela no Word
                num_rows = tabela["num_linhas"]
                num_cols = tabela["num_colunas"]
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'

                for row_idx, linha in enumerate(tabela["linhas"]):
                    for col_idx, celula in enumerate(linha):
                        if col_idx < num_cols:
                            table.cell(row_idx, col_idx).text = str(celula)

                doc.add_paragraph()  # Espaço após tabela

        doc.save(filepath)
        logger.info(f"DOCX exportado: {filepath}")
        return str(filepath)


class OCRProcessor:
    """Classe principal para processar documentos com OCR."""

    def __init__(self):
        self.ocr = GoogleVisionOCR()
        self.exporter = DocumentExporter()

    def process_document(
        self,
        file_path: str,
        export_format: str = "auto",
        generate_summary: bool = True
    ) -> Dict[str, Any]:
        """
        Processa um documento completo com OCR.

        Args:
            file_path: Caminho do arquivo
            export_format: 'csv', 'xlsx', 'docx', ou 'auto' (detecta automaticamente)
            generate_summary: Se deve gerar resumo

        Returns:
            Dict com resultado do processamento e caminhos dos arquivos exportados
        """
        # Processa com OCR
        resultado = self.ocr.process_image(file_path)

        # Gera resumo
        if generate_summary:
            resultado["resumo"] = self.ocr.generate_summary(resultado["texto_completo"])

        # Define formato de exportação
        if export_format == "auto":
            if resultado["tabelas"]:
                export_format = "xlsx"
            else:
                export_format = "docx"

        # Nome base do arquivo
        filename = Path(file_path).stem

        # Exporta
        resultado["arquivos_exportados"] = {}

        if export_format in ["csv", "xlsx"] and resultado["tabelas"]:
            if export_format == "csv":
                resultado["arquivos_exportados"]["csv"] = self.exporter.export_to_csv(
                    resultado["tabelas"], filename
                )
            else:
                resultado["arquivos_exportados"]["xlsx"] = self.exporter.export_to_xlsx(
                    resultado["tabelas"], filename
                )

        # Sempre gera DOCX com texto completo
        resultado["arquivos_exportados"]["docx"] = self.exporter.export_to_docx(
            resultado, filename, incluir_resumo=generate_summary
        )

        return resultado

    def process_comprovante(self, file_path: str, parcela_id: int) -> Dict[str, Any]:
        """
        Processa um comprovante de pagamento específico.

        Args:
            file_path: Caminho do arquivo do comprovante
            parcela_id: ID da parcela relacionada

        Returns:
            Dict com informações extraídas do comprovante
        """
        resultado = self.ocr.process_image(file_path)
        resultado["resumo"] = self.ocr.generate_summary(resultado["texto_completo"])

        # Extrai informações específicas de comprovante
        info = self.ocr._extract_key_info(resultado["texto_completo"])

        resultado["comprovante_info"] = {
            "parcela_id": parcela_id,
            "valor_detectado": info.get("Valor Principal"),
            "data_detectada": info.get("Data"),
            "tipo_documento": self.ocr._identify_document_type(resultado["texto_completo"]),
            "cpf_cnpj_detectado": info.get("CPF") or info.get("CNPJ"),
            "nome_detectado": info.get("Nome")
        }

        return resultado


# Instância global para uso nos endpoints
ocr_processor = OCRProcessor()
